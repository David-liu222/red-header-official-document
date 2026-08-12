#!/usr/bin/env python3
"""Apply the current red-header rules to an existing DOCX.

The source is read-only. The output is a new DOCX. This is intentionally
conservative: it adds the red-header block and normalizes formatting without
rewriting the document's substantive text.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


RED = "FF0000"
BODY_FONT = "仿宋"
TITLE_FONT = "黑体"
# The approved 66号 .doc has no direct rFonts on the red-header runs. Its
# effective CJK paragraph default exports as 宋体, so preserve that template
# behavior instead of substituting a conventional but unproven font.
HEADER_FONT = "宋体"


def set_east_asia_font(run, name: str):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)


def set_run(run, *, font: str, size: float, bold: bool | None, color: str | None = None):
    set_east_asia_font(run, font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def remove_paragraph(paragraph):
    paragraph._p.getparent().remove(paragraph._p)


def replace_text(paragraph, text: str):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    return run


def set_spacing(paragraph, *, indent_pt: float | None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(28)
    pf.first_line_indent = Pt(indent_pt) if indent_pt is not None else None


def set_single_spacing(paragraph):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0


def clear_title_before_after_spacing(paragraph):
    """Match the approved template's inherited 0-line before/after spacing.

    WPS displays a title with no direct before/after spacing as “0 行”. Writing
    point-based zero or explicit beforeLines/afterLines zero instead displays
    “0 磅” or “自动”, which does not match the supplied template.
    """
    ppr = paragraph._p.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is not None:
        for attribute in ("before", "after", "beforeLines", "afterLines", "beforeAutospacing", "afterAutospacing"):
            spacing.attrib.pop(qn(f"w:{attribute}"), None)


def set_exact_spacing(paragraph, points: float):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(points)


def set_bottom_border(paragraph, color=RED, size="16"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = pbdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pbdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def add_before(anchor, text, *, size, font, bold, color=None, border=False, line_spacing=None):
    paragraph = anchor._parent.add_paragraph()
    anchor._p.addprevious(paragraph._p)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if line_spacing is None:
        set_single_spacing(paragraph)
    else:
        set_exact_spacing(paragraph, line_spacing)
    pf.keep_with_next = True
    run = paragraph.add_run(text)
    set_run(run, font=font, size=size, bold=bold, color=color)
    if border:
        set_bottom_border(paragraph)
    return paragraph


def normalize_signature(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(paragraph, indent_pt=None)
    for run in paragraph.runs:
        set_run(run, font=BODY_FONT, size=14, bold=False)


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(paragraph)
    run = paragraph.add_run()
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])


def apply(
    input_path: Path,
    output_path: Path,
    *,
    title: str,
    issuing_number: str,
    title_continuations: list[str] | None = None,
    no_recipient: bool = False,
):
    doc = Document(str(input_path))
    if not doc.paragraphs:
        raise ValueError("文档没有正文段落，无法安全套用红头。")

    original_paragraphs = list(doc.paragraphs)
    title_paragraph = next((p for p in original_paragraphs if p.text.strip() == title), None)
    if title_paragraph is None:
        raise ValueError(f"未找到标题段落：{title}")
    title_paragraphs = [title_paragraph]
    for continuation in title_continuations or []:
        paragraph = next((p for p in original_paragraphs if p.text.strip() == continuation), None)
        if paragraph is None:
            raise ValueError(f"未找到标题续行段落：{continuation}")
        title_paragraphs.append(paragraph)

    # A leading blank paragraph from a legacy DOC must not push the newly
    # inserted red head down the page. Remove only whitespace-only paragraphs
    # that occur before the first title line; never remove content elsewhere.
    for paragraph in original_paragraphs:
        if paragraph is title_paragraph:
            break
        if not paragraph.text.strip():
            remove_paragraph(paragraph)

    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(23)
    section.bottom_margin = Mm(23)
    section.left_margin = Mm(23)
    section.right_margin = Mm(23)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(15)
    add_page_number(section)

    # Add red-head paragraphs before the original title, in reverse insertion order.
    add_before(title_paragraph, "内 蒙 古 准 格 尔 旗", size=45, font=HEADER_FONT, bold=True, color=RED, line_spacing=65)
    add_before(title_paragraph, "力量煤业有限公司文件", size=45, font=HEADER_FONT, bold=True, color=RED, line_spacing=65)
    code_text = issuing_number if issuing_number else "〔发文字号待补〕"
    add_before(title_paragraph, code_text, size=14, font=BODY_FONT, bold=False, color=None, border=True, line_spacing=27)
    add_before(title_paragraph, "", size=10, font=BODY_FONT, bold=False, line_spacing=10)

    # Title: 黑体二号, normal weight, centered; inherit template 0-line
    # before/after spacing so WPS shows “0 行”.
    # The approved 24-number template uses Normal (正文文本) rather than the
    # source document's Heading 3 style. Heading 3 makes WPS show 自动 here.
    for paragraph in title_paragraphs:
        paragraph.style = doc.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_single_spacing(paragraph)
        clear_title_before_after_spacing(paragraph)
        paragraph.paragraph_format.first_line_indent = None
        for run in paragraph.runs:
            set_run(run, font=TITLE_FONT, size=22, bold=False)

    recipient = None if no_recipient else next(
        (p for p in original_paragraphs if p not in title_paragraphs and p.text.strip()),
        None,
    )
    signature_markers = ("内蒙古准格尔旗力量煤业有限公司", "法定代表人", "日期：")
    for paragraph in original_paragraphs:
        text = paragraph.text.strip()
        if paragraph in title_paragraphs:
            continue
        if any(marker in text for marker in signature_markers) or re.fullmatch(r"\d{4}年\d{1,2}月\d{1,2}日", text):
            normalize_signature(paragraph)
            continue
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        indent = None if paragraph is recipient else 28
        set_spacing(paragraph, indent_pt=indent)
        for run in paragraph.runs:
            set_run(run, font=BODY_FONT, size=14, bold=False)

    # Main recipient is top-aligned; keep its text intact and remove accidental leading spaces.
    if recipient is not None:
        text = recipient.text.strip()
        if text != recipient.text:
            replace_text(recipient, text)
        recipient.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(recipient, indent_pt=None)
        for run in recipient.runs:
            set_run(run, font=BODY_FONT, size=14, bold=True)

    doc.core_properties.title = title
    doc.core_properties.subject = "套红头公文"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main():
    parser = argparse.ArgumentParser(description="为现有 DOCX 加套红头并按固定规则排版")
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--title", required=True)
    parser.add_argument("--title-continuation", action="append", default=[], help="标题的续行文字，可重复传入")
    parser.add_argument("--no-recipient", action="store_true", help="原稿无主送单位时使用，避免把首段正文误当主送单位")
    parser.add_argument("--issuing-number", default="〔发文字号待补〕")
    args = parser.parse_args()
    if args.input.resolve() == args.output.resolve():
        raise SystemExit("禁止覆盖源文件，请指定新的输出路径。")
    apply(
        args.input,
        args.output,
        title=args.title,
        issuing_number=args.issuing_number,
        title_continuations=args.title_continuation,
        no_recipient=args.no_recipient,
    )
    print(args.output)


if __name__ == "__main__":
    main()
